import os
from pathlib import Path
from typing import Optional
from docx import Document as DocxDocument
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from fpdf import FPDF
from odf.opendocument import OpenDocumentText, load as load_odt
from odf.text import P as OdfParagraph
from odf import text as odf_text
from striprtf.striprtf import rtf_to_text
from bs4 import BeautifulSoup
import markdown

try:
    pass
except ImportError:
    print("Install required packages: pip install -r requirements.txt")


class DocumentConverter:
    """Convert between various document formats."""
    
    SUPPORTED_READ_FORMATS = {'.txt', '.docx', '.pdf', '.rtf', '.odt', '.html', '.htm', '.md'}
    SUPPORTED_WRITE_FORMATS = {'.txt', '.docx', '.pdf', '.rtf', '.odt', '.html', '.md'}
    SUPPORTED_FORMATS = SUPPORTED_READ_FORMATS | SUPPORTED_WRITE_FORMATS
    
    def __init__(self):
        self.output_dir = Path.home() / "Desktop" / "Converted Documents"
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def convert(self, input_path: str, output_format: str) -> Optional[str]:
        """Convert a document to the specified format."""
        input_file = Path(input_path)
        
        if not input_file.exists():
            raise FileNotFoundError(f"File not found: {input_path}")
        
        input_format = input_file.suffix.lower()
        output_format = output_format.lower()
        
        if not output_format.startswith('.'):
            output_format = f'.{output_format}'
        
        if input_format not in self.SUPPORTED_READ_FORMATS:
            raise ValueError(f"Cannot read format: {input_format}")
        
        if output_format not in self.SUPPORTED_WRITE_FORMATS:
            raise ValueError(f"Cannot write format: {output_format}")
        
        # Read content based on input format
        content = self._read_file(input_file, input_format)
        
        # Write content based on output format
        output_path = self.output_dir / f"{input_file.stem}{output_format}"
        self._write_file(output_path, content, output_format)
        
        return str(output_path)
    
    def _read_file(self, file_path: Path, file_format: str) -> str:
        """Read document content."""
        if file_format == '.txt':
            return file_path.read_text(encoding='utf-8')
        elif file_format == '.docx':
            doc = DocxDocument(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif file_format == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_format == '.rtf':
            return self._read_rtf(file_path)
        elif file_format == '.odt':
            return self._read_odt(file_path)
        elif file_format in ('.html', '.htm'):
            return self._read_html(file_path)
        elif file_format == '.md':
            return file_path.read_text(encoding='utf-8')
        else:
            raise ValueError(f"Cannot read format: {file_format}")
    
    def _read_rtf(self, file_path: Path) -> str:
        """Read RTF file content."""
        rtf_content = file_path.read_text(encoding='utf-8', errors='ignore')
        return rtf_to_text(rtf_content)
    
    def _read_odt(self, file_path: Path) -> str:
        """Read ODT file content."""
        doc = load_odt(str(file_path))
        paragraphs = doc.getElementsByType(odf_text.P)
        return '\n'.join([str(p) for p in paragraphs])
    
    def _read_html(self, file_path: Path) -> str:
        """Read HTML file and extract text content."""
        html_content = file_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        return soup.get_text(separator='\n', strip=True)
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using OCR."""
        images = convert_from_path(file_path)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image) + "\n"
        return text
    
    def _write_file(self, output_path: Path, content: str, file_format: str) -> None:
        """Write content to file."""
        if file_format == '.txt':
            output_path.write_text(content, encoding='utf-8')
        elif file_format == '.docx':
            doc = DocxDocument()
            for paragraph in content.split('\n'):
                doc.add_paragraph(paragraph)
            doc.save(output_path)
        elif file_format == '.pdf':
            self._write_pdf(output_path, content)
        elif file_format == '.rtf':
            self._write_rtf(output_path, content)
        elif file_format == '.odt':
            self._write_odt(output_path, content)
        elif file_format == '.html':
            self._write_html(output_path, content)
        elif file_format == '.md':
            output_path.write_text(content, encoding='utf-8')
        else:
            raise ValueError(f"Cannot write format: {file_format}")
    
    def _write_pdf(self, output_path: Path, content: str) -> None:
        """Write content to PDF file."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", size=12)
        # Handle unicode by encoding
        for line in content.split('\n'):
            pdf.multi_cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'))
        pdf.output(str(output_path))
    
    def _write_rtf(self, output_path: Path, content: str) -> None:
        """Write content to RTF file."""
        # Basic RTF structure
        rtf_content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}\f0\fs24 "
        for line in content.split('\n'):
            # Escape special RTF characters
            escaped = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            rtf_content += escaped + r"\par "
        rtf_content += "}"
        output_path.write_text(rtf_content, encoding='utf-8')
    
    def _write_odt(self, output_path: Path, content: str) -> None:
        """Write content to ODT file."""
        doc = OpenDocumentText()
        for line in content.split('\n'):
            p = OdfParagraph(text=line)
            doc.text.addElement(p)
        doc.save(str(output_path))
    
    def _write_html(self, output_path: Path, content: str) -> None:
        """Write content to HTML file."""
        paragraphs = '\n'.join([f'    <p>{line}</p>' for line in content.split('\n') if line.strip()])
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
</head>
<body>
{paragraphs}
</body>
</html>"""
        output_path.write_text(html_content, encoding='utf-8')


if __name__ == "__main__":
    converter = DocumentConverter()
    # Example: converter.convert("sample.pdf", ".txt")